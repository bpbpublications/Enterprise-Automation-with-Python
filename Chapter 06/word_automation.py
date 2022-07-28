from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading("My Automation Document", 0)

p = document.add_paragraph("Autommating Processes")
p.add_run("bold").bold = True

document.add_heading("Process Automation", level=1)
document.add_paragraph("Automation is the first step towards growth", style="Intense Quote")

document.add_paragraph("Automate PDF", style="List Bullet")
document.add_paragraph("Automate Word", style="List Number")

document.add_page_break()

document.save("test_doc.docx")
